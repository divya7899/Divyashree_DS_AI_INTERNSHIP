# Create a Professional Project Documentation Report
# Format similar to the uploaded Project_Code_Documentation.docx

# Create a Professional Project Documentation Report
# Format similar to the uploaded Project_Code_Documentation.docx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from datetime import datetime

# Create document
document = Document()

# ================================
# TITLE
# ================================
document.add_heading("Customer Analytics Project: Exploratory Data Analysis Documentation", level=1)

document.add_paragraph(
    "This document provides the technical explanation and structured documentation "
    "for the Customer Analytics Exploratory Data Analysis (EDA) project. "
    "It outlines the workflow, implementation logic, analysis methodology, and key findings."
)

# ================================
# TABLE OF CONTENTS (Manual Section Listing)
# ================================
document.add_heading("Table of Contents", level=2)
document.add_paragraph("1. Project Overview")
document.add_paragraph("2. Environment & Libraries Used")
document.add_paragraph("3. Data Loading & Initial Inspection")
document.add_paragraph("4. Data Cleaning & Preprocessing")
document.add_paragraph("5. Univariate Analysis")
document.add_paragraph("6. Bivariate & Multivariate Analysis")
document.add_paragraph("7. Correlation Analysis")
document.add_paragraph("8. Key Insights & Executive Summary")

# ================================
# 1. PROJECT OVERVIEW
# ================================
document.add_heading("1. Project Overview", level=2)
document.add_paragraph(
    "The objective of this project is to perform structured Exploratory Data Analysis (EDA) "
    "on a customer analytics dataset. Each row in the dataset represents a single customer, "
    "including demographic attributes (Age, Gender, City, etc.) and behavioral metrics "
    "(Annual Income, Spending Score, Purchase Activity, etc.)."
)

document.add_paragraph(
    "The primary goal of this analysis is to identify patterns, relationships, "
    "data quality issues, and potential business insights using statistical and visual techniques."
)

# ================================
# 2. ENVIRONMENT
# ================================
document.add_heading("2. Environment & Libraries Used", level=2)
document.add_paragraph("The following Python libraries were used in this project:")
document.add_paragraph("- Pandas: Data manipulation and analysis")
document.add_paragraph("- NumPy: Numerical computations")
document.add_paragraph("- Matplotlib: Data visualization")
document.add_paragraph("- Seaborn: Advanced statistical visualization")

# ================================
# 3. DATA LOADING
# ================================
document.add_heading("3. Data Loading & Initial Inspection", level=2)
document.add_paragraph(
    "The dataset was loaded using Pandas read_csv() function. "
    "Initial inspection was performed using the following methods:"
)
document.add_paragraph("- head(): To preview the first few records")
document.add_paragraph("- info(): To examine data types and missing values")
document.add_paragraph("- describe(): To generate statistical summaries")
document.add_paragraph(
    "This step helped in understanding dataset dimensions, feature types, and potential inconsistencies."
)

# ================================
# 4. DATA CLEANING
# ================================
document.add_heading("4. Data Cleaning & Preprocessing", level=2)
document.add_paragraph(
    "Data preprocessing was performed to ensure reliability and analytical accuracy."
)

document.add_paragraph(
    "• Missing Values: Missing numerical values were handled using mean imputation. "
    "Categorical missing values were filled using mode or retained based on relevance."
)

document.add_paragraph(
    "• Duplicate Records: Duplicate rows were identified using duplicated() and removed "
    "to prevent biased statistical analysis."
)

document.add_paragraph(
    "• Data Type Verification: Columns were validated to ensure correct numeric or categorical types."
)

# ================================
# 5. UNIVARIATE
# ================================
document.add_heading("5. Univariate Analysis", level=2)
document.add_paragraph(
    "Univariate analysis was performed to understand the distribution of individual variables."
)
document.add_paragraph(
    "- Histograms were used to visualize numerical distributions such as Age and Annual Income."
)
document.add_paragraph(
    "- Bar charts were used for categorical variables such as Gender and Education."
)
document.add_paragraph(
    "This analysis revealed the concentration patterns and variability within each feature."
)

# ================================
# 6. BIVARIATE
# ================================
document.add_heading("6. Bivariate & Multivariate Analysis", level=2)
document.add_paragraph(
    "Bivariate analysis was conducted to examine relationships between two variables."
)
document.add_paragraph(
    "- Scatter plots were used to analyze the relationship between Annual Income and Spending Score."
)
document.add_paragraph(
    "- Boxplots were used to compare Spending Score across different genders."
)
document.add_paragraph(
    "These visualizations helped in identifying trends and behavioral differences across customer segments."
)

# ================================
# 7. CORRELATION
# ================================
document.add_heading("7. Correlation Analysis", level=2)
document.add_paragraph(
    "A correlation matrix was generated using numeric features to identify multi-variable relationships."
)
document.add_paragraph(
    "The heatmap visualization highlighted moderate correlations between income and spending behavior, "
    "indicating potential predictive relationships."
)

# ================================
# 8. EXECUTIVE SUMMARY
# ================================
document.add_heading("8. Key Insights & Executive Summary", level=2)
document.add_paragraph(
    "1. Younger customers demonstrate relatively higher spending behavior compared to older groups."
)
document.add_paragraph(
    "2. Annual income shows measurable influence on spending score patterns."
)
document.add_paragraph(
    "3. Gender does not exhibit a strong statistical impact on overall spending distribution."
)
document.add_paragraph(
    "4. Data cleaning was essential to remove noise such as duplicates and missing values before analysis."
)

document.add_paragraph(
    "This structured EDA provides a strong foundation for further steps such as customer segmentation, "
    "predictive modeling, or marketing strategy optimization."
)

# Save document
file_path = "/mnt/data/Customer_Analytics_Project_Documentation.docx"
document.save(file_path)

file_path

